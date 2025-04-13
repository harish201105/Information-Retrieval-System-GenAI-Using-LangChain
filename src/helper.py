import os
import tempfile
import time
from typing import List, Dict, Any

from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.docstore.document import Document as LangchainDocument
from langchain.chains.summarize import load_summarize_chain
from google.api_core.exceptions import ResourceExhausted, InvalidArgument
from dotenv import load_dotenv


load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")  
os.environ['GOOGLE_API_KEY'] = GOOGLE_API_KEY

# Maximum retries for API quota errors
MAX_RETRIES = 3
RETRY_DELAY = 5  # seconds


def check_api_key():
    """Verify that the API key is valid and not expired"""
    if not GOOGLE_API_KEY or GOOGLE_API_KEY.strip() == "":
        raise ValueError("Google API key is missing. Please add it to your .env file as GOOGLE_API_KEY=your_key_here")
    
    # Try a minimal API call to verify the key works
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", task_type="retrieval_document")
        _ = embeddings.embed_query("Test query to verify API key")
        return True
    except Exception as e:
        error_msg = str(e)
        if "API_KEY_INVALID" in error_msg or "expired" in error_msg:
            raise ValueError(f"Google API key is invalid or expired. Please renew your API key. Error: {error_msg}")
        elif "PERMISSION_DENIED" in error_msg:
            raise ValueError(f"Google API key doesn't have permission to access this service. Error: {error_msg}")
        else:
            raise ValueError(f"Error validating Google API key: {error_msg}")


def get_file_text(file):
    """Extract text from various file types"""
    text = ""
    file_extension = os.path.splitext(file.name)[1].lower()
    
    # Create a temporary file to save the uploaded file
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
        temp_file.write(file.getvalue())
        temp_path = temp_file.name
    
    try:
        if file_extension == ".pdf":
            # Handle PDF files
            pdf_reader = PdfReader(temp_path)
            for page in pdf_reader.pages:
                text += page.extract_text()
        elif file_extension == ".docx":
            # Handle DOCX files
            doc = Document(temp_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif file_extension == ".txt":
            # Handle TXT files
            with open(temp_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            text = f"Unsupported file format: {file_extension}"
    finally:
        # Clean up the temporary file
        if os.path.exists(temp_path):
            os.remove(temp_path)
    
    return text


def get_pdf_text(files):
    """Process multiple files and extract text"""
    text = ""
    for file in files:
        text += get_file_text(file) + "\n\n"
    return text


def get_text_chunks(text):
    """Split text into manageable chunks"""
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    chunks = text_splitter.split_text(text)
    return chunks


def get_vector_store(text_chunks):
    """Create a vector store from text chunks"""
    # First verify the API key
    try:
        check_api_key()
    except ValueError as e:
        raise Exception(f"API Key Error: {str(e)}")
        
    # Retry logic for API quota errors
    retries = 0
    last_error = None
    
    while retries < MAX_RETRIES:
        try:
            embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", task_type="retrieval_document")
            vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
            return vector_store
        except ResourceExhausted as e:
            last_error = e
            retries += 1
            if retries < MAX_RETRIES:
                print(f"API quota exceeded. Retrying in {RETRY_DELAY} seconds... (Attempt {retries}/{MAX_RETRIES})")
                time.sleep(RETRY_DELAY)
            else:
                raise Exception(f"Failed after {MAX_RETRIES} attempts due to API quota limits. Please try again later or reduce the document size.") from e
        except InvalidArgument as e:
            error_msg = str(e)
            if "API_KEY_INVALID" in error_msg or "expired" in error_msg:
                raise Exception(f"Google API key is invalid or expired. Please renew your API key in the .env file.")
            else:
                raise Exception(f"Error with Google API: {error_msg}")
        except Exception as e:
            error_msg = str(e)
            if "API key" in error_msg.lower() or "invalid" in error_msg.lower() or "expired" in error_msg.lower():
                raise Exception(f"Google API key issue: {error_msg}. Please check your .env file and update your API key.")
            else:
                raise Exception(f"Error creating vector store: {error_msg}")


def get_conversational_chain(vector_store):
    """Create a conversational chain with source documents"""
    # First verify the API key
    try:
        check_api_key()
    except ValueError as e:
        raise Exception(f"API Key Error: {str(e)}")
    
    # Initialize memory for conversation history
    memory = ConversationBufferMemory(
        memory_key="chat_history",
        return_messages=True,
        output_key="answer"  # Specify which output key to store in memory
    )
    
    # Create a retriever from the vector store
    retriever = vector_store.as_retriever(
        search_type="similarity",
        search_kwargs={"k": 5}  # Return top 5 most similar chunks
    )
    
    try:
        # Create the conversational chain
        llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro", temperature=0.2)
        
        chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=retriever,
            memory=memory,
            return_source_documents=True,
            output_key="answer"  # Specify the output key
        )
        
        return chain
    except Exception as e:
        error_msg = str(e)
        if "API key" in error_msg.lower() or "invalid" in error_msg.lower() or "expired" in error_msg.lower():
            raise Exception(f"Google API key issue: {error_msg}. Please check your .env file and update your API key.")
        else:
            raise Exception(f"Error creating conversation chain: {error_msg}")


def get_document_summary(text):
    """Generate a summary of the document with efficient chunking and fallback"""
    # First verify the API key
    try:
        check_api_key()
    except ValueError as e:
        return fallback_extractive_summary(text, error_message=str(e))
    
    # If text is too long, use a more efficient approach
    if len(text) > 50000:  # For very long documents
        # Create a simple extractive summary by taking the first 500 chars of each document
        # and the first 3 sentences of each paragraph
        paragraphs = text.split('\n\n')
        summary_parts = []
        
        # Take first paragraph as is (usually contains important context)
        if paragraphs:
            summary_parts.append(paragraphs[0])
        
        # For other paragraphs, take first 3 sentences or first 200 chars
        for i, para in enumerate(paragraphs[1:10]):  # Limit to first 10 paragraphs for efficiency
            if para.strip():
                if i == 0:
                    # Include full first paragraph
                    summary_parts.append(para)
                else:
                    # For other paragraphs, take first 3 sentences
                    sentences = para.split('.')
                    summary_parts.append('.'.join(sentences[:3]) + '.')
        
        # Try to use AI to refine this extractive summary
        extracted_summary = '\n\n'.join(summary_parts)
        return refine_summary(extracted_summary)
    
    # For shorter documents, use the standard approach with retry logic
    return standard_summary(text)


def standard_summary(text):
    """Generate summary using LangChain's summarization chain"""
    retries = 0
    last_error = None
    
    while retries < MAX_RETRIES:
        try:
            llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro", temperature=0.3)
            
            # Convert text to LangChain documents with larger chunks to reduce API calls
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=15000, chunk_overlap=1500)
            docs = [LangchainDocument(page_content=t) for t in text_splitter.split_text(text)]
            
            # If we have many chunks, just use the first few
            if len(docs) > 5:
                docs = docs[:5]
            
            # Use map_reduce for multiple documents
            chain = load_summarize_chain(llm, chain_type="map_reduce")
            summary = chain.run(docs)
            
            return summary
        except ResourceExhausted as e:
            last_error = e
            retries += 1
            if retries < MAX_RETRIES:
                print(f"API quota exceeded. Retrying in {RETRY_DELAY} seconds... (Attempt {retries}/{MAX_RETRIES})")
                time.sleep(RETRY_DELAY)
            else:
                # Fallback to extractive summary if API limits are reached
                print(f"API quota limit reached after {MAX_RETRIES} attempts. Using fallback extractive summary.")
                return fallback_extractive_summary(text, error_message=str(e))
        except Exception as e:
            # For any other error, fall back to extractive summary
            error_msg = str(e)
            if "API key" in error_msg.lower() or "invalid" in error_msg.lower() or "expired" in error_msg.lower():
                return fallback_extractive_summary(text, error_message=f"API key issue: {error_msg}")
            else:
                last_error = e
                print(f"Error generating summary: {error_msg}")
                return fallback_extractive_summary(text, error_message=error_msg)


def refine_summary(extracted_text):
    """Refine an extractive summary using the LLM"""
    try:
        llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro", temperature=0.3)
        prompt = f"Please refine and improve the following extracted document summary:\n\n{extracted_text}\n\nProvide a coherent, well-structured summary:"
        return llm.invoke(prompt).content
    except ResourceExhausted as e:
        print(f"API quota limit reached when refining summary: {str(e)}")
        return fallback_extractive_summary(extracted_text, error_message=str(e))
    except Exception as e:
        error_msg = str(e)
        if "API key" in error_msg.lower() or "invalid" in error_msg.lower() or "expired" in error_msg.lower():
            return fallback_extractive_summary(extracted_text, error_message=f"API key issue: {error_msg}")
        else:
            print(f"Error refining summary: {error_msg}")
            return fallback_extractive_summary(extracted_text, error_message=error_msg)


def fallback_extractive_summary(text, error_message=None):
    """Create a simple extractive summary when API calls fail"""
    # Split text into paragraphs and sentences
    paragraphs = text.split('\n\n')
    summary_parts = []
    
    # Take first paragraph as is (usually contains important context)
    if paragraphs and paragraphs[0].strip():
        summary_parts.append(paragraphs[0])
    
    # For other paragraphs, take first 3 sentences or first 200 chars from important paragraphs
    # Try to find paragraphs that seem important (contain keywords like summary, conclusion, etc.)
    important_keywords = ['summary', 'conclusion', 'result', 'finding', 'important', 'key', 'significant']
    
    # First pass: Look for paragraphs with important keywords
    important_paragraphs = []
    for para in paragraphs[1:]:
        if para.strip():
            para_lower = para.lower()
            if any(keyword in para_lower for keyword in important_keywords):
                important_paragraphs.append(para)
    
    # Add important paragraphs first (up to 3)
    for para in important_paragraphs[:3]:
        sentences = para.split('.')
        summary_parts.append('.'.join(sentences[:3]) + '.')
    
    # If we don't have enough content yet, add regular paragraphs
    regular_paragraphs = [p for p in paragraphs[1:10] if p.strip() and p not in important_paragraphs]
    for para in regular_paragraphs[:5 - len(important_paragraphs)]:
        sentences = para.split('.')
        summary_parts.append('.'.join(sentences[:3]) + '.')
    
    # Format the summary with clear section breaks
    formatted_summary = '\n\n'.join(summary_parts)
    
    # Add a note explaining this is a fallback summary with more context
    fallback_note = """

[Note: This is an extractive summary created due to API limitations. It contains key excerpts from the document but has not been refined by AI.]"""
    
    # If there's a specific error message, include it in a more user-friendly way
    if error_message:
        if "API key" in error_message.lower() or "invalid" in error_message.lower() or "expired" in error_message.lower():
            fallback_note += """

Error: Google API key is invalid or has expired. Please update your API key in the .env file.

To fix this issue:
1. Get a new API key from Google AI Studio (https://ai.google.dev/)
2. Update the GOOGLE_API_KEY value in your .env file
3. Restart the application"""
        elif "429" in error_message or "quota" in error_message.lower():
            fallback_note += """

Error: Google API quota limit reached. Please try again later or reduce document size.

To get better results:
- Try again in a few minutes when API quotas reset
- Upload a smaller document
- Split your document into smaller parts"""
        else:
            fallback_note += f"""

Error: {error_message}"""
    
    return formatted_summary + fallback_note